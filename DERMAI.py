import io
import matplotlib.pyplot as plt
import numpy as np
from tensorflow.keras.models import load_model
from PIL import Image
from docx import Document
from docx.shared import Inches
import time

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Load the pretrained Keras model
model = load_model(r"C:\Users\rupin\Downloads\VIT VELLORE HACKATHON\skinn.h5")

# Define class labels
class_labels = [
    "Acne and Rosacea Photos",
    "Actinic Keratosis Basal Cell Carcinoma and other Malignant Lesions",
    "Atopic Dermatitis Photos",
    "Bullous Disease Photos",
    "Cellulitis Impetigo and other Bacterial Infections",
    "Eczema Photos",
    "Exanthems and Drug Eruptions",
    "Hair Loss Photos Alopecia and other Hair Diseases",
    "Herpes HPV and other STDs Photos",
    "Light Diseases and Disorders of Pigmentation",
    "Lupus and other Connective Tissue diseases",
    "Melanoma Skin Cancer Nevi and Moles",
    "Nail Fungus and other Nail Disease",
    "Poison Ivy Photos and other Contact Dermatitis",
    "Psoriasis pictures Lichen Planus and related diseases",
    "Scabies Lyme Disease and other Infestations and Bites",
    "Seborrheic Keratoses and other Benign Tumors",
    "Systemic Disease",
    "Tinea Ringworm Candidiasis and other Fungal Infections",
    "Urticaria Hives",
    "Vascular Tumors",
    "Vasculitis Photos",
    "Warts Molluscum and other Viral Infections"
]

# Define descriptions and treatment methods for each disease
# Define descriptions and treatment methods for each disease
disease_info = {
    'Acne and Rosacea Photos': {
        'description': 'Acne is a common skin condition that causes pimples and redness on the face. '
                       'Rosacea is a chronic skin condition that causes redness and visible blood vessels.',
        'treatment': 'Treatment for acne may include topical creams, antibiotics, and lifestyle changes. '
                     'Rosacea treatment may involve medications and lifestyle adjustments.'
    },
    'Actinic Keratosis Basal Cell Carcinoma and other Malignant Lesions': {
        'description': 'Actinic keratosis is a precancerous skin condition, while basal cell carcinoma and malignant lesions '
                       'are types of skin cancer.',
        'treatment': 'Treatment for actinic keratosis may include cryotherapy or topical creams. Basal cell carcinoma and '
                     'malignant lesion treatment may involve surgery or radiation.'
    },
    'Atopic Dermatitis Photos': {
        'description': 'Atopic dermatitis, also known as eczema, is a chronic skin condition that causes itching and inflammation.',
        'treatment': 'Eczema treatment includes moisturizers, topical steroids, and avoiding triggers that worsen symptoms.'
    },
    'Bullous Disease Photos': {
        'description': 'Bullous diseases are a group of rare skin conditions characterized by the development of large fluid-filled blisters.',
        'treatment': 'Treatment depends on the specific type of bullous disease and may involve medications, wound care, or other interventions.'
    },
    'Cellulitis Impetigo and other Bacterial Infections': {
        'description': 'These conditions are bacterial skin infections that can cause redness, swelling, and discomfort.',
        'treatment': 'Treatment includes antibiotics and proper wound care, depending on the specific infection.'
    },
    'Exanthems and Drug Eruptions': {
        'description': 'Exanthems are skin rashes often caused by viral infections, while drug eruptions are adverse reactions to medications.',
        'treatment': 'Treatment varies depending on the cause, and may involve managing the underlying condition or discontinuing the medication.'
    },
    'Hair Loss Photos Alopecia and other Hair Diseases': {
        'description': 'Conditions affecting the hair and scalp, including hair loss (alopecia) and various hair disorders.',
        'treatment': 'Treatment depends on the specific condition and may involve medications, lifestyle changes, or other interventions.'
    },
    'Herpes HPV and other STDs Photos': {
        'description': 'Sexually transmitted diseases (STDs) such as herpes and human papillomavirus (HPV) can cause skin symptoms.',
        'treatment': 'Treatment depends on the specific STD and may involve antiviral medications or other therapies.'
    },
    'Light Diseases and Disorders of Pigmentation': {
        'description': 'Conditions affecting skin pigmentation, including light-related disorders.',
        'treatment': 'Treatment varies based on the specific condition and may include phototherapy or topical treatments.'
    },
    'Lupus and other Connective Tissue diseases': {
        'description': 'Connective tissue diseases like lupus can affect the skin and other organs.',
        'treatment': 'Treatment is tailored to the individual and may involve medications to manage symptoms and control the underlying disease.'
    },
    'Melanoma Skin Cancer Nevi and Moles': {
        'description': 'Melanoma is a serious form of skin cancer that originates in pigment-producing cells. Nevi are common moles on the skin.',
        'treatment': 'Melanoma treatment may involve surgery, chemotherapy, radiation, or immunotherapy. Moles are typically monitored for changes.'
    },
    'Nail Fungus and other Nail Disease': {
        'description': 'Nail fungus and other nail diseases can cause changes in the appearance and health of the nails.',
        'treatment': 'Treatment for nail fungus may include antifungal medications, while other nail diseases require specific medical care.'
    },
    'Poison Ivy Photos and other Contact Dermatitis': {
        'description': 'Poison ivy and other plants can cause contact dermatitis, leading to itchy rashes.',
        'treatment': 'Treatment includes avoiding contact with the irritant, over-the-counter creams, or prescription medications.'
    },
    'Psoriasis pictures Lichen Planus and related diseases': {
        'description': 'Psoriasis is an autoimmune skin condition that causes red, scaly patches. Lichen planus is an inflammatory condition affecting the skin, mouth, and genitals.',
        'treatment': 'Psoriasis treatment includes topical treatments, phototherapy, and medications. Lichen planus treatment depends on the affected area.'
    },
    'Scabies Lyme Disease and other Infestations and Bites': {
        'description': 'These conditions involve infestations and bites by parasites or insects.',
        'treatment': 'Treatment depends on the specific infestation or bite and may include medications and proper hygiene.'
    },
    'Seborrheic Keratoses and other Benign Tumors': {
        'description': 'Seborrheic keratoses are noncancerous skin growths. Benign tumors are noncancerous lumps.',
        'treatment': 'Seborrheic keratosis removal may be done for cosmetic reasons. Benign tumors may not require treatment unless causing symptoms.'
    },
    'Systemic Disease': {
        'description': 'Systemic diseases can affect various organs, including the skin.',
        'treatment': 'Treatment depends on the specific systemic disease and may involve medications and management of the underlying condition.'
    },
    'Tinea Ringworm Candidiasis and other Fungal Infections': {
        'description': 'These are fungal skin infections that can affect different areas of the body.',
        'treatment': 'Treatment for fungal infections typically includes antifungal medications or creams and good hygiene practices.'
    },
    'Urticaria Hives': {
        'description': 'Urticaria, commonly known as hives, is an allergic skin reaction that causes red, itchy welts on the skin.',
        'treatment': 'Urticaria treatment includes antihistamines and identifying and avoiding triggers.'
    },
    'Vascular Tumors': {
        'description': 'Vascular tumors are growths that involve blood vessels. They can be benign or malignant.',
        'treatment': 'Treatment for vascular tumors depends on their type and location.'
    },
    'Vasculitis Photos': {
        'description': 'Vasculitis is an inflammation of blood vessels that can affect various organs, including the skin.',
        'treatment': 'Treatment for vasculitis may include medications to reduce inflammation and manage symptoms.'
    },
    'Warts Molluscum and other Viral Infections': {
        'description': 'Warts and molluscum are viral skin infections that can cause raised bumps on the skin.',
        'treatment': 'Treatment for warts and molluscum may involve topical treatments or other removal methods.'
    }
}

# Rest of the code (predict_skin_disease, send_diagnosis
def predict_skin_disease(image_path):
    img = Image.open(image_path)
    img = img.resize((224, 224))
    img = np.array(img)
    img = img / 255.0
    img = np.expand_dims(img, axis=0)
    predictions = model.predict(img)
    predicted_class_index = np.argmax(predictions)
    predicted_class = class_labels[predicted_class_index]
    return predicted_class, predictions.max(), img[0]

# Function to send the diagnosis report via email
def send_diagnosis_report(patient_name, patient_email, diagnosis_report, app_password, attachment_path=None):
    # Email configurations
    smtp_server = "smtp.gmail.com"
    smtp_port = 587  # For Gmail
    sender_email = "medilink0000@gmail.com"  # Your Gmail address
    recipient_email = patient_email
    subject = "Diagnosis Report for " + patient_name

    # Create a message container
    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = recipient_email
    message["Subject"] = subject

    # Attach the diagnosis report
    message.attach(MIMEText(diagnosis_report, "plain"))

    if attachment_path:
        # Attach the Word file
        with open(attachment_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename=diagnosis_report.docx")
            message.attach(part)

    try:
        # Connect to the SMTP server
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()

        # Log in with the App Password
        server.login(sender_email, app_password)

        # Send the email
        server.sendmail(sender_email, recipient_email, message.as_string())

        # Quit the server
        server.quit()

        print(f"Email sent to {patient_name} at {patient_email}")

    except Exception as e:
        print(f"Failed to send email: {str(e)}")

# Function to perform diagnosis and send report
def perform_diagnosis_and_send_report(image_path, patient_name, patient_email, app_password):
    # Perform the skin disease prediction and get information
    predicted_class, confidence, input_image = predict_skin_disease(image_path)
    info = disease_info.get(predicted_class, {})

    # Create a figure to display the image and the predicted disease
    plt.figure()
    plt.imshow(input_image)
    plt.title("Input Image - Predicted: " + predicted_class)
    plt.axis('off')

    # Create a timestamp to make the document name unique
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    document_name = f"Skin_Disease_Prediction_{timestamp}.docx"

    # Save the result to a Word document
    doc = Document()
    doc.add_heading("Skin Disease Prediction - DermAI", 0)

    doc.add_heading("Predicted Disease:", level=1)
    doc.add_paragraph(predicted_class)

    doc.add_heading("Accuracy:", level=1)
    doc.add_paragraph(f"{confidence * 100:.2f}%")

    input_image_path = "input_image.jpg"
    input_image = (input_image * 255).astype(np.uint8)
    Image.fromarray(input_image).save(input_image_path)

    doc.add_heading("Input Image:", level=1)
    doc.add_picture(input_image_path, width=Inches(4.5))

    # Include description and treatment information if available
    if info:
        doc.add_heading("Disease Description:", level=1)
        doc.add_paragraph(info.get('description', 'Description not available'))

        doc.add_heading("Treatment Methods:", level=1)
        doc.add_paragraph(info.get('treatment', 'Treatment information not available'))

    doc.save(document_name)

    # Send the diagnosis report via email
    send_diagnosis_report(patient_name, patient_email, "Diagnosis Report", app_password, document_name)

    # Show the input image and the predicted disease
    plt.show()

    print(f"Results saved to '{document_name}' and email sent to {patient_name} at {patient_email}")

app_password = "lpjd ginj ifij qqjz"  # Replace with your actual App Password
image_path = input("Enter the path to the image: ")
patient_name = input("Enter Patient Name:")
patient_email = input("Enter Patient Email Id:")
perform_diagnosis_and_send_report(image_path, patient_name, patient_email, app_password)
